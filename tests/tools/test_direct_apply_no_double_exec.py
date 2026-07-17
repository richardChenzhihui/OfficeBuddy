"""Regression tests for the old preview/apply double-execution bug class.

Each non-idempotent operation is dispatched exactly once through the real tool
layer and must affect the document exactly once.
"""
import pytest
from docx import Document
from openpyxl import load_workbook

from office_agent.tools import REGISTRY, ToolContext


@pytest.fixture
def ctx():
    c = ToolContext()
    yield c
    c.sessions.close_all()


def _open(ctx, path):
    result = REGISTRY.dispatch(ctx, "open_document", {"file_path": str(path)})
    assert result["success"], result
    return result["doc_id"]


def test_word_append_applies_exactly_once(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    session = ctx.sessions.get(doc_id)
    before = session.doc.paragraphs[0].text
    result = REGISTRY.dispatch(
        ctx,
        "word_edit_text",
        {
            "doc_id": doc_id,
            "selector": {"type": "paragraph", "index": 0},
            "operation": "append",
            "text": "_SUFFIX",
        },
    )
    assert result["success"], result
    after = session.doc.paragraphs[0].text
    assert after == before + "_SUFFIX"
    assert after.count("_SUFFIX") == 1  # not doubled


def test_word_insert_element_applies_exactly_once(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    session = ctx.sessions.get(doc_id)
    n_before = len(session.doc.paragraphs)
    result = REGISTRY.dispatch(
        ctx,
        "word_insert_element",
        {"doc_id": doc_id, "element_type": "paragraph", "content": "UNIQUE_MARKER_XYZ"},
    )
    assert result["success"], result
    texts = [p.text for p in session.doc.paragraphs]
    assert texts.count("UNIQUE_MARKER_XYZ") == 1
    assert len(session.doc.paragraphs) == n_before + 1


def test_excel_insert_rows_applies_exactly_once(ctx, excel_doc_path):
    doc_id = _open(ctx, excel_doc_path)
    session = ctx.sessions.get(doc_id)
    sheet = session.doc.sheetnames[0]
    ws = session.doc[sheet]
    a1_before = ws["A1"].value
    max_row_before = ws.max_row
    result = REGISTRY.dispatch(
        ctx,
        "excel_insert_rows_cols",
        {"doc_id": doc_id, "sheet": sheet, "position": 1, "count": 1, "insert_type": "row"},
    )
    assert result["success"], result
    assert ws.max_row == max_row_before + 1  # exactly one row inserted
    assert ws["A2"].value == a1_before  # shifted exactly once


def test_excel_create_chart_applies_exactly_once(ctx, excel_doc_path):
    doc_id = _open(ctx, excel_doc_path)
    session = ctx.sessions.get(doc_id)
    sheet = session.doc.sheetnames[0]
    result = REGISTRY.dispatch(
        ctx,
        "excel_create_chart",
        {"doc_id": doc_id, "sheet": sheet, "data_range": "A1:B2", "chart_type": "bar"},
    )
    assert result["success"], result
    assert len(session.doc[sheet]._charts) == 1  # exactly one chart


def test_excel_append_fill_mode_applies_exactly_once(ctx, excel_doc_path):
    doc_id = _open(ctx, excel_doc_path)
    session = ctx.sessions.get(doc_id)
    sheet = session.doc.sheetnames[0]
    ws = session.doc[sheet]
    before = str(ws["A1"].value or "")
    result = REGISTRY.dispatch(
        ctx,
        "excel_write_cells",
        {
            "doc_id": doc_id,
            "sheet": sheet,
            "range": "A1",
            "values": [["_X"]],
            "fill_mode": "append",
        },
    )
    assert result["success"], result
    assert str(ws["A1"].value) == before + "_X"


def test_mutating_call_creates_snapshot_and_undo_works(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    session = ctx.sessions.get(doc_id)
    before = session.doc.paragraphs[0].text
    result = REGISTRY.dispatch(
        ctx,
        "word_edit_text",
        {
            "doc_id": doc_id,
            "selector": {"type": "paragraph", "index": 0},
            "operation": "replace",
            "text": "REPLACED",
        },
    )
    assert "snapshot_id" in result
    undo_result = REGISTRY.dispatch(ctx, "undo", {"doc_id": doc_id, "steps": 1})
    assert undo_result["success"], undo_result
    assert ctx.sessions.get(doc_id).doc.paragraphs[0].text == before


def test_selector_no_match_reports_actionable_error(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    result = REGISTRY.dispatch(
        ctx,
        "word_edit_text",
        {
            "doc_id": doc_id,
            "selector": {"type": "text_match", "contains": "NONEXISTENT_ZEBRA"},
            "operation": "replace",
            "text": "x",
        },
    )
    assert result["success"] is False
    assert "matched no elements" in result["error"]
    assert "paragraphs" in result["error"]  # includes document stats


def test_unknown_tool_reports_available_tools(ctx):
    result = REGISTRY.dispatch(ctx, "nonexistent_tool", {})
    assert result["success"] is False
    assert "available tools" in result["error"]


def test_invalid_args_report_validation_details(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    result = REGISTRY.dispatch(
        ctx, "word_edit_text", {"doc_id": doc_id, "selector": {"type": "paragraph"}}
    )
    assert result["success"] is False
    assert "Invalid arguments" in result["error"]


def test_wrong_doc_type_is_actionable(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    result = REGISTRY.dispatch(
        ctx, "excel_read_cells", {"doc_id": doc_id, "sheet": "Sheet1"}
    )
    assert result["success"] is False
    assert "word_* tools" in result["error"]


def test_save_document_defaults_to_edited_copy(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    original_bytes = word_doc_path.read_bytes()
    REGISTRY.dispatch(
        ctx,
        "word_edit_text",
        {
            "doc_id": doc_id,
            "selector": {"type": "paragraph", "index": 0},
            "operation": "replace",
            "text": "SAVED_EDIT",
        },
    )
    result = REGISTRY.dispatch(ctx, "save_document", {"doc_id": doc_id})
    assert result["success"], result
    assert result["saved_path"].endswith(".edited.docx")
    assert word_doc_path.read_bytes() == original_bytes  # original untouched
    assert Document(result["saved_path"]).paragraphs[0].text == "SAVED_EDIT"


def test_save_over_original_without_overwrite_is_refused(ctx, word_doc_path):
    doc_id = _open(ctx, word_doc_path)
    result = REGISTRY.dispatch(
        ctx, "save_document", {"doc_id": doc_id, "path": str(word_doc_path)}
    )
    assert result["success"] is False
    assert "Refusing to overwrite" in result["error"]


def test_tool_schemas_are_valid_anthropic_shape(ctx):
    tools = REGISTRY.to_anthropic_tools()
    assert len(tools) >= 13
    for tool in tools:
        assert set(tool) == {"name", "description", "input_schema"}
        assert tool["input_schema"]["type"] == "object"
        assert tool["description"]
