import pytest
from docx import Document

from office_agent.core.selector_parser import SelectorError, SelectorParser


@pytest.fixture
def doc():
    d = Document()
    d.add_paragraph("Alpha introduction", style="Heading 1")
    d.add_paragraph("Body text about apples")
    d.add_paragraph("More BODY text about oranges")
    d.add_table(rows=2, cols=2)
    return d


def test_paragraph_by_index(doc):
    result = SelectorParser.parse_word_selector({"type": "paragraph", "index": 1}, doc)
    assert len(result) == 1
    assert result[0].text == "Body text about apples"


def test_paragraph_index_out_of_range_is_actionable(doc):
    with pytest.raises(SelectorError, match="has 3 paragraphs"):
        SelectorParser.parse_word_selector({"type": "paragraph", "index": 99}, doc)


def test_paragraph_range(doc):
    result = SelectorParser.parse_word_selector(
        {"type": "paragraph", "range": (0, 1)}, doc
    )
    assert [p.text for p in result] == ["Alpha introduction", "Body text about apples"]


def test_text_match_contains_first(doc):
    result = SelectorParser.parse_word_selector(
        {"type": "text_match", "contains": "body", "case_sensitive": False}, doc
    )
    assert len(result) == 1
    assert result[0].text == "Body text about apples"


def test_text_match_all_occurrences(doc):
    result = SelectorParser.parse_word_selector(
        {
            "type": "text_match",
            "contains": "body",
            "case_sensitive": False,
            "occurrence": "all",
        },
        doc,
    )
    assert len(result) == 2


def test_text_match_regex(doc):
    result = SelectorParser.parse_word_selector(
        {"type": "text_match", "regex": r"apples|oranges", "occurrence": "all"}, doc
    )
    assert len(result) == 2


def test_text_match_no_hit_returns_empty(doc):
    result = SelectorParser.parse_word_selector(
        {"type": "text_match", "contains": "zebra"}, doc
    )
    assert result == []


def test_style_match(doc):
    result = SelectorParser.parse_word_selector(
        {"type": "style_match", "style_name": "Heading 1"}, doc
    )
    assert len(result) == 1
    assert result[0].text == "Alpha introduction"


def test_table_cell_drill_down(doc):
    result = SelectorParser.parse_word_selector(
        {"type": "table", "table_index": 0, "row_index": 1, "cell_index": 0}, doc
    )
    assert len(result) == 1


def test_table_row_out_of_range(doc):
    with pytest.raises(SelectorError, match="has 2 row"):
        SelectorParser.parse_word_selector(
            {"type": "table", "table_index": 0, "row_index": 9}, doc
        )


def test_header_selector_returns_paragraphs(doc):
    result = SelectorParser.parse_word_selector({"type": "header"}, doc)
    assert isinstance(result, list)


def test_header_section_out_of_range(doc):
    with pytest.raises(SelectorError, match="section"):
        SelectorParser.parse_word_selector({"type": "header", "section_index": 5}, doc)


def test_excel_missing_sheet_is_actionable():
    from openpyxl import Workbook

    wb = Workbook()
    with pytest.raises(SelectorError, match="not found"):
        SelectorParser.parse_excel_selector({"sheet": "Nope"}, wb)


def test_excel_valid_range():
    from openpyxl import Workbook

    wb = Workbook()
    ws, coords = SelectorParser.parse_excel_selector(
        {"sheet": "Sheet", "range": "A1:B2"}, wb
    )
    assert coords == (1, 1, 2, 2)


def test_excel_bad_range_is_actionable():
    from openpyxl import Workbook

    wb = Workbook()
    with pytest.raises(SelectorError, match="Invalid Excel range"):
        SelectorParser.parse_excel_selector({"sheet": "Sheet", "range": "!!!"}, wb)
