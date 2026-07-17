"""Word adapter wrapping python-docx"""
from typing import Any, Optional, List, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from ..schemas.operations import StyleParams, TextOperation
from ..schemas.selector import WordSelector


class WordAdapter:
    """Adapter for python-docx operations"""
    
    @staticmethod
    def read_content(elements: List[Any]) -> List[Dict[str, Any]]:
        """Read content from document elements"""
        result = []
        
        for element in elements:
            if hasattr(element, 'text'):  # Paragraph or Run
                result.append({
                    "type": "text",
                    "text": element.text,
                    "style": element.style.name if hasattr(element, 'style') else None
                })
            elif hasattr(element, 'rows'):  # Table
                table_data = []
                for row in element.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                result.append({
                    "type": "table",
                    "data": table_data
                })
        
        return result
    
    @staticmethod
    def edit_text(elements: List[Any], operation: TextOperation, 
                 text: Optional[str] = None) -> Dict[str, Any]:
        """Edit text in elements"""
        affected = []
        
        for element in elements:
            if hasattr(element, 'text'):
                old_text = element.text
                
                if operation == TextOperation.REPLACE:
                    element.text = text or ""
                    affected.append(f"Replaced: '{old_text[:50]}...' -> '{text[:50] if text else ''}...'")
                
                elif operation == TextOperation.INSERT:
                    if text:
                        # For paragraphs, use add_run to preserve formatting
                        if hasattr(element, 'runs'):
                            # Insert at beginning: clear existing runs and add new text first
                            element.text = text
                            # Note: This will lose existing formatting, but INSERT at beginning is less common
                            # For better formatting preservation, we'd need to insert a run at position 0
                        else:
                            # For runs, prepend text
                            element.text = text + element.text
                        affected.append(f"Inserted: '{text[:50]}...'")
                
                elif operation == TextOperation.DELETE:
                    element.text = ""
                    affected.append(f"Deleted: '{old_text[:50]}...'")
                
                elif operation == TextOperation.APPEND:
                    if text:
                        # For paragraphs, use add_run to preserve existing formatting
                        if hasattr(element, 'runs'):
                            # Get the last run's formatting to maintain style consistency
                            if element.runs:
                                last_run = element.runs[-1]
                                new_run = element.add_run(text)
                                # Copy formatting from last run to maintain consistency
                                new_run.font.name = last_run.font.name
                                if last_run.font.size:
                                    new_run.font.size = last_run.font.size
                                new_run.font.bold = last_run.font.bold
                                new_run.font.italic = last_run.font.italic
                                new_run.font.underline = last_run.font.underline
                                if last_run.font.color and last_run.font.color.rgb:
                                    new_run.font.color.rgb = last_run.font.color.rgb
                            else:
                                # No existing runs, just add text
                                element.add_run(text)
                        else:
                            # For runs, append directly
                            element.text = element.text + text
                        affected.append(f"Appended: '{text[:50]}...'")
        
        return {
            "affected": affected,
            "operation": operation.value
        }
    
    @staticmethod
    def apply_style(elements: List[Any], style: StyleParams) -> Dict[str, Any]:
        """Apply style to elements"""
        affected = []
        
        for element in elements:
            if hasattr(element, 'runs'):  # Paragraph
                for run in element.runs:
                    WordAdapter._apply_run_style(run, style)
                affected.append("Paragraph styled")
            
            elif hasattr(element, 'font'):  # Run
                WordAdapter._apply_run_style(element, style)
                affected.append("Run styled")
            
            elif hasattr(element, 'paragraph_format'):  # Paragraph
                para_format = element.paragraph_format
                if style.paragraph_spacing:
                    para_format.space_before = Pt(style.paragraph_spacing)
                if style.line_spacing:
                    para_format.line_spacing = style.line_spacing
                if style.alignment:
                    alignment_map = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    para_format.alignment = alignment_map.get(style.alignment, WD_ALIGN_PARAGRAPH.LEFT)
                affected.append("Paragraph format applied")
        
        return {
            "affected": affected,
            "style_applied": {
                "font": bool(style.font_name or style.font_size or style.bold or style.italic),
                "color": bool(style.color),
                "alignment": bool(style.alignment)
            }
        }
    
    @staticmethod
    def _apply_run_style(run: Any, style: StyleParams):
        """Apply style to a run"""
        font = run.font
        
        if style.font_name:
            font.name = style.font_name
        if style.font_size:
            font.size = Pt(style.font_size)
        if style.bold is not None:
            font.bold = style.bold
        if style.italic is not None:
            font.italic = style.italic
        if style.underline is not None:
            # True = single underline (WD_UNDERLINE.SINGLE), False = no underline (WD_UNDERLINE.NONE)
            font.underline = style.underline
        
        if style.color:
            # Parse color (hex or RGB)
            if style.color.startswith('#'):
                rgb = RGBColor.from_string(style.color[1:])
            else:
                # Assume RGB tuple or hex without #
                try:
                    rgb = RGBColor.from_string(style.color)
                except:
                    rgb = RGBColor(0, 0, 0)  # Default to black
            font.color.rgb = rgb
    
    @staticmethod
    def insert_element(doc: Document, position: Optional[int], 
                      element_type: str, content: Any) -> Dict[str, Any]:
        """Insert element into document"""
        if element_type == "table":
            # content should be List[List[str]]
            if not isinstance(content, list):
                return {"error": "Table content must be a list of lists"}
            
            if position is not None and position < len(doc.paragraphs):
                para = doc.paragraphs[position]
                table = doc.add_table(rows=len(content), cols=len(content[0]) if content else 0)
                for i, row_data in enumerate(content):
                    for j, cell_data in enumerate(row_data):
                        table.rows[i].cells[j].text = str(cell_data)
                para._element.addnext(table._element)
            else:
                table = doc.add_table(rows=len(content), cols=len(content[0]) if content else 0)
                for i, row_data in enumerate(content):
                    for j, cell_data in enumerate(row_data):
                        table.rows[i].cells[j].text = str(cell_data)
            
            return {
                "element_type": "table",
                "rows": len(content),
                "cols": len(content[0]) if content else 0
            }
        
        elif element_type == "paragraph":
            if position is not None:
                para = doc.paragraphs[position].insert_paragraph_before(str(content))
            else:
                para = doc.add_paragraph(str(content))
            
            return {
                "element_type": "paragraph",
                "text": str(content)[:50]
            }
        
        elif element_type == "page_break":
            if position is not None and position < len(doc.paragraphs):
                para = doc.paragraphs[position]
                run = para.add_run()
                run.add_break(6)  # Page break
            else:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_break(6)
            
            return {
                "element_type": "page_break"
            }
        
        else:
            return {"error": f"Unsupported element type: {element_type}"}
    
    @staticmethod
    def find_replace(doc: Document, find: str, replace: str, 
                    scope: Optional[List[Any]] = None) -> Dict[str, Any]:
        """Find and replace text"""
        count = 0
        affected = []
        
        elements = scope if scope else doc.paragraphs
        
        for element in elements:
            if hasattr(element, 'text') and find in element.text:
                element.text = element.text.replace(find, replace)
                count += element.text.count(replace)
                affected.append(f"Element: '{element.text[:50]}...'")
        
        return {
            "count": count,
            "affected": affected,
            "find": find,
            "replace": replace
        }
