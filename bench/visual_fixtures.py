"""Fixtures for the visual-trap battery (M1 / M3).

Each fixture is deliberately *clean*: its own trap detector must report
`present=False` on the freshly built file. The trap only springs if an agent
executes the task's instruction the naive way — producing a file that is
schema-valid, that byte-level assertions call correct, and that is visibly
broken once real Word/Excel lays it out.

Design constraint for every fixture here: the instruction must be a thing a
normal user would actually ask for. A trap you can only trigger by being
absurd measures nothing.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Twips
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

# python-docx's default template is US Letter with 1.25" side margins:
# 12240 - 1800 - 1800 = 8640 twips of usable text width. (Measured, not
# assumed — the docs commonly quote 9360 for 1" margins, which is not what
# this template ships.)
TEXT_WIDTH_TWIPS = 8640


# --------------------------------------------------------------------------
# VT1 — CJK font slot
# --------------------------------------------------------------------------
def vis_cjk(path: Path) -> None:
    """Chinese body text with no explicit font anywhere.

    Clean because nothing is pinned: Word resolves CJK through the theme's
    minorEastAsia font and every glyph renders in the same face. The trap is
    an agent that sets only the Latin slots (rFonts/ascii+hAnsi) when asked to
    change the font — Word then has to guess a fallback face per character and
    the paragraph renders visibly uneven.
    """
    d = Document()
    d.add_heading("城市轨道交通运营年报", level=1)
    d.add_paragraph(
        "本年度全网运营里程达到 1,284 公里，较上年增长 8.6%，日均客运量 1,120 万人次。"
    )
    d.add_paragraph(
        "全年共开通 3 条新线，新增车站 47 座，其中换乘站 12 座。设备可靠度指标全部达标。"
    )
    d.add_paragraph(
        "客运服务方面，全网列车正点率 99.87%，乘客有效投诉率同比下降 21%，满意度调查得分 92.4 分。"
    )
    d.add_paragraph(
        "下一年度将重点推进智慧车站建设，计划完成 30 座车站的自动化改造并试点全自动运行线路。"
    )
    d.save(str(path))


# --------------------------------------------------------------------------
# VT2 — Excel column overflow (########)
# --------------------------------------------------------------------------
def vis_narrow(path: Path) -> None:
    """Amounts sit in a column that is only just wide enough for the raw
    numbers. Clean as built (General format renders short). Applying a
    thousands-separator currency format without widening the column makes
    Excel render the whole column as ########.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "费用台账"
    ws.append(["科目", "所属部门", "金额"])
    rows = [
        ["办公用品", "行政部", 1284630.5],
        ["差旅费", "销售部", 2375180.75],
        ["市场推广", "市场部", 4860220.0],
        ["设备采购", "技术部", 7392455.25],
        ["培训费", "人力资源部", 1045990.4],
    ]
    for r in rows:
        ws.append(r)
    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 14
    # 9 chars: fits "1284630.5" (General) but not "1,284,630.50" (12 chars).
    ws.column_dimensions["C"].width = 9
    for cell in ws[1]:
        cell.font = Font(bold=True)
    wb.save(str(path))


# --------------------------------------------------------------------------
# VT3 — Word table wider than the text column
# --------------------------------------------------------------------------
def vis_table(path: Path) -> None:
    """A 3-column table that already spans essentially the full text width.

    Clean as built (9300 of 9360 twips). Appending a fourth column without
    re-dividing the existing widths pushes the table past the right margin,
    and real Word clips the overflow off the page edge.
    """
    d = Document()
    d.add_heading("季度供应商结算表", level=1)
    d.add_paragraph("下表列出本季度已完成结算的供应商及其结算金额。")

    table = d.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    col_w = Twips(2860)  # 3 x 2860 = 8580, just inside the 8640 text width
    header = ["供应商名称", "结算金额（元）", "结算日期"]
    rows = [
        ["华信物流有限公司", "1,284,630", "2026-04-12"],
        ["恒通电子科技", "876,200", "2026-04-27"],
        ["明远办公服务", "342,880", "2026-05-08"],
        ["中辰建筑工程", "2,015,440", "2026-05-30"],
    ]
    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text
    for ri, row in enumerate(rows, start=1):
        for ci, text in enumerate(row):
            table.rows[ri].cells[ci].text = text
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_w

    d.add_paragraph("结算金额以财务系统记录为准。")
    d.save(str(path))


# --------------------------------------------------------------------------
# VT4 — chart anchored on top of its own data
# --------------------------------------------------------------------------
def vis_chart(path: Path) -> None:
    """Plain data table, no chart yet. The trap is anchoring a new chart
    inside the used range so the plot area covers the numbers it is plotting.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "月度销量"
    ws.append(["月份", "华东", "华北"])
    data = [
        ["1月", 320, 244], ["2月", 298, 261], ["3月", 415, 302],
        ["4月", 388, 295], ["5月", 442, 331], ["6月", 470, 358],
    ]
    for r in data:
        ws.append(r)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    wb.save(str(path))


# --------------------------------------------------------------------------
# VT5 — invisible text (white on a light fill)
# --------------------------------------------------------------------------
def vis_contrast(path: Path) -> None:
    """Header row already carries a *light* fill. Clean as built: dark text on
    a light background. Setting the header font to white — a completely
    ordinary request, and the right move on a dark fill — makes the header
    text invisible unless the agent also darkens the fill.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "项目进度"
    ws.append(["项目名称", "负责人", "完成度", "截止日期"])
    rows = [
        ["智慧园区一期", "张伟", "78%", "2026-09-30"],
        ["数据中台重构", "李娜", "45%", "2026-11-15"],
        ["移动端改版", "王强", "92%", "2026-08-20"],
        ["风控模型升级", "刘敏", "31%", "2026-12-31"],
    ]
    for r in rows:
        ws.append(r)
    light = PatternFill("solid", fgColor="FFF2CC")  # light amber
    for cell in ws[1]:
        cell.fill = light
        cell.font = Font(bold=True, color="333333")
        cell.alignment = Alignment(horizontal="center")
    for i, w in enumerate([22, 12, 10, 14], start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    wb.save(str(path))


# --------------------------------------------------------------------------
# VT6 — orphaned heading at the foot of a page
# --------------------------------------------------------------------------
# Tuned so "市场展望" sits low on page 1 but still has body text under it.
# Enlarging the heading pushes that body text to page 2 and strands the
# heading alone at the bottom. Word decides pagination, not us, so this value
# is measured rather than reasoned about — re-derive it with:
#   python selftest_traps.py --calibrate-vt6
ORPHAN_FILLER_PARAGRAPHS = 6

ORPHAN_HEADING = "市场展望"


def vis_orphan(path: Path, filler: int = ORPHAN_FILLER_PARAGRAPHS) -> None:
    """The section heading is *manually formatted* (bold + larger run on a
    normal paragraph), not a real Heading style — extremely common in
    hand-made documents, and crucially it carries no keepNext, so Word will
    happily leave it stranded at the bottom of a page.
    """
    d = Document()
    title = d.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("年度经营分析报告")
    tr.bold = True
    tr.font.size = Pt(18)

    # Deliberately does not repeat the heading text: the render-side detector
    # matches whole text blocks, and a body sentence echoing the heading would
    # make "which block is the heading" ambiguous.
    d.add_paragraph("本报告分为经营回顾与后市研判两个部分，供管理层决策参考。")

    sub = d.add_paragraph()
    sr = sub.add_run("经营回顾")
    sr.bold = True
    sr.font.size = Pt(14)

    for i in range(1, filler + 1):
        d.add_paragraph(
            f"{i}. 本年度公司在核心业务板块保持稳健增长，营业收入较上年同期提升，"
            f"毛利率维持在合理区间。渠道结构持续优化，直营与经销的比例趋于均衡，"
            f"库存周转天数下降，应收账款账龄结构改善。"
        )

    heading = d.add_paragraph()
    hr = heading.add_run(ORPHAN_HEADING)
    hr.bold = True
    hr.font.size = Pt(14)

    d.add_paragraph(
        "展望下一年度，行业整体需求预计温和复苏，公司将围绕产品升级与渠道下沉两条主线布局。"
    )
    d.add_paragraph(
        "同时公司将加大海外市场投入，目标是三年内海外收入占比提升至 30% 以上。"
    )
    d.save(str(path))


# --------------------------------------------------------------------------
# VT7 — wrapped text clipped by a fixed row height
# --------------------------------------------------------------------------
def vis_rowclip(path: Path) -> None:
    """Long remarks in a narrow column, wrap OFF and rows at single-line
    height. Clean as built: with wrap off the text simply spills into the empty
    cells to the right and stays readable. Turning wrap on without raising the
    row height clips every remark to its first line.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "验收记录"
    ws.append(["批次", "备注"])
    remarks = [
        "外观检查合格，包装完整无破损，随附质检报告与合格证，建议按标准流程入库。",
        "第二批次存在轻微色差，已与供应商确认属工艺允差范围内，可正常验收。",
        "运输过程中外箱受潮，开箱抽检未见产品受影响，建议后续更换防潮包装。",
        "数量与订单一致，但到货时间晚于约定三天，已按合同条款记录并知会采购。",
    ]
    for i, r in enumerate(remarks, start=1):
        ws.append([f"B{i:03d}", r])
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 28
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for row in range(1, len(remarks) + 2):
        ws.row_dimensions[row].height = 15  # single line
    wb.save(str(path))


# --------------------------------------------------------------------------
# VT8 — merging the header row destroys the column labels
# --------------------------------------------------------------------------
MERGE_HEADERS = ["产品", "销量", "单价", "金额"]


def vis_merge(path: Path) -> None:
    """A normal table whose first row carries the column labels. Adding a
    spanning title by merging *that* row silently discards every label but the
    first — openpyxl and Excel both keep only the top-left value. The file is
    valid, the title is there, and the table has lost its header.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "销售汇总"
    ws.append(MERGE_HEADERS)
    for row in [
        ["A1 智能音箱", 1280, 299, 382720],
        ["B2 无线耳机", 2140, 199, 425860],
        ["C3 智能手环", 960, 399, 383040],
        ["D4 便携音箱", 1530, 259, 396270],
    ]:
        ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for i, w in enumerate([18, 10, 10, 12], start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    wb.save(str(path))


# --------------------------------------------------------------------------
# VT9 — table torn across printed pages (render-only)
# --------------------------------------------------------------------------
WIDE_HEADERS = ["订单号", "客户名称", "下单日期", "金额", "状态", "备注"]


def vis_wide(path: Path) -> None:
    """Six columns that currently fit inside one printed page.

    Clean as built. Widening the remarks column — the obvious way to make its
    content readable — pushes the sheet past one page wide, and Excel prints
    the overflowing columns onto a separate sheet of paper. On screen it looks
    fine; the printed/exported artefact is a table torn in half. Nothing in the
    file says so: column widths are all perfectly valid.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "订单明细"
    ws.append(WIDE_HEADERS)
    rows = [
        ["SO-20260101", "远见科技有限公司", "2026-01-04", 128400, "已发货",
         "客户要求分两批交付，第一批已于当月完成，第二批待确认排产计划。"],
        ["SO-20260102", "启明制造集团", "2026-01-09", 264800, "生产中",
         "定制规格，需提供第三方检测报告后方可安排出库。"],
        ["SO-20260103", "海通物流股份", "2026-01-15", 87600, "已完成",
         "常规订单，按标准条款执行，无特殊要求。"],
        ["SO-20260104", "长风新材料", "2026-01-22", 341200, "待付款",
         "客户信用额度已用满，需财务审批后释放订单。"],
        ["SO-20260105", "同辉电子", "2026-01-28", 156300, "已发货",
         "含赠品两件，随货同行单已注明，需回收签收回执。"],
    ]
    for r in rows:
        ws.append(r)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    # Total ~70 chars — inside the ~78 chars a portrait Letter page prints.
    for col, w in zip("ABCDEF", [14, 18, 12, 10, 8, 12]):
        ws.column_dimensions[col].width = w
    wb.save(str(path))


BUILDERS = {
    "vis_cjk.docx": vis_cjk,
    "vis_narrow.xlsx": vis_narrow,
    "vis_table.docx": vis_table,
    "vis_chart.xlsx": vis_chart,
    "vis_contrast.xlsx": vis_contrast,
    "vis_orphan.docx": vis_orphan,
    "vis_rowclip.xlsx": vis_rowclip,
    "vis_merge.xlsx": vis_merge,
    "vis_wide.xlsx": vis_wide,
}


def register() -> None:
    """Make these fixtures reachable through the existing fixtures.build()."""
    import fixtures

    fixtures.BUILDERS.update(BUILDERS)


# --------------------------------------------------------------------------
# Naive edits: what a byte-level agent plausibly does. Used by the selftest to
# prove each trap actually springs — a trap that never fires measures nothing.
# --------------------------------------------------------------------------
def naive_cjk(path: Path) -> None:
    """Set the font the python-docx way: Latin slots only."""
    d = Document(str(path))
    for para in d.paragraphs:
        for run in para.runs:
            run.font.name = "宋体"
            run.font.size = Pt(12)
    d.save(str(path))


def naive_narrow(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["费用台账"]
    for row in ws.iter_rows(min_row=2, min_col=3, max_col=3):
        for cell in row:
            cell.number_format = "#,##0.00"
    wb.save(str(path))


def naive_table(path: Path) -> None:
    """Append a 4th column without re-dividing the existing widths."""
    d = Document(str(path))
    table = d.tables[0]
    table.add_column(Twips(2400))
    table.rows[0].cells[3].text = "备注"
    for r in range(1, len(table.rows)):
        table.rows[r].cells[3].text = "已对账"
    d.save(str(path))


def naive_chart(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["月度销量"]
    chart = BarChart()
    chart.type = "col"
    data = Reference(ws, min_col=2, max_col=3, min_row=1, max_row=7)
    cats = Reference(ws, min_col=1, min_row=2, max_row=7)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    ws.add_chart(chart, "B2")  # squarely on top of the data
    wb.save(str(path))


def naive_contrast(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["项目进度"]
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
    wb.save(str(path))


def naive_orphan(path: Path) -> None:
    d = Document(str(path))
    for para in d.paragraphs:
        if para.text.strip() == ORPHAN_HEADING:
            for run in para.runs:
                run.font.size = Pt(22)
    d.save(str(path))


def naive_rowclip(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["验收记录"]
    for row in ws.iter_rows(min_row=2, min_col=2, max_col=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    wb.save(str(path))


def naive_merge(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["销售汇总"]
    ws.merge_cells("A1:D1")  # eats the column labels
    ws["A1"] = "2026 年第一季度销售汇总"
    ws["A1"].alignment = Alignment(horizontal="center")
    wb.save(str(path))


def naive_wide(path: Path) -> None:
    """Widen the remarks column until the text fits — and the page does not."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    wb["订单明细"].column_dimensions["F"].width = 60
    wb.save(str(path))


NAIVE_EDITS = {
    "vis_cjk.docx": naive_cjk,
    "vis_narrow.xlsx": naive_narrow,
    "vis_table.docx": naive_table,
    "vis_chart.xlsx": naive_chart,
    "vis_contrast.xlsx": naive_contrast,
    "vis_orphan.docx": naive_orphan,
    "vis_rowclip.xlsx": naive_rowclip,
    "vis_merge.xlsx": naive_merge,
    "vis_wide.xlsx": naive_wide,
}


# --------------------------------------------------------------------------
# Correct edits: what a system that *noticed* would deliver. Used by the
# selftest to prove the detector is not simply always-on.
# --------------------------------------------------------------------------
def correct_cjk(path: Path) -> None:
    from docx.oxml.ns import qn

    d = Document(str(path))
    for para in d.paragraphs:
        for run in para.runs:
            run.font.name = "宋体"
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    d.save(str(path))


def correct_narrow(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["费用台账"]
    for row in ws.iter_rows(min_row=2, min_col=3, max_col=3):
        for cell in row:
            cell.number_format = "#,##0.00"
    ws.column_dimensions["C"].width = 16  # widened to fit the formatted value
    wb.save(str(path))


def correct_table(path: Path) -> None:
    from docx.oxml.ns import qn

    d = Document(str(path))
    table = d.tables[0]
    share = Twips(TEXT_WIDTH_TWIPS // 4 - 20)  # 2140 x 4 = 8560, fits
    table.add_column(share)
    table.rows[0].cells[3].text = "备注"
    for r in range(1, len(table.rows)):
        table.rows[r].cells[3].text = "已对账"
    for row in table.rows:  # re-divide so the total still fits the text width
        for cell in row.cells:
            cell.width = share
    for gc in table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol")):
        gc.set(qn("w:w"), str(share.twips))
    d.save(str(path))


def correct_chart(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["月度销量"]
    chart = BarChart()
    chart.type = "col"
    data = Reference(ws, min_col=2, max_col=3, min_row=1, max_row=7)
    cats = Reference(ws, min_col=1, min_row=2, max_row=7)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    ws.add_chart(chart, "E2")  # clear of the used range
    wb.save(str(path))


def correct_contrast(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["项目进度"]
    dark = PatternFill("solid", fgColor="1F4E79")
    for cell in ws[1]:
        cell.fill = dark  # darken the fill so white text is legible
        cell.font = Font(bold=True, color="FFFFFF")
    wb.save(str(path))


def correct_orphan(path: Path) -> None:
    """Enlarge the heading but keep it with the text that follows."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = Document(str(path))
    for para in d.paragraphs:
        if para.text.strip() == ORPHAN_HEADING:
            for run in para.runs:
                run.font.size = Pt(22)
            pPr = para._element.get_or_add_pPr()
            keep = OxmlElement("w:keepNext")
            pPr.append(keep)
    d.save(str(path))


def correct_rowclip(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["验收记录"]
    for row in ws.iter_rows(min_row=2, min_col=2, max_col=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            ws.row_dimensions[cell.row].height = 45  # room for the wrapped lines
    wb.save(str(path))


def correct_merge(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["销售汇总"]
    ws.insert_rows(1)  # give the title its own row; labels stay intact
    ws.merge_cells("A1:D1")
    ws["A1"] = "2026 年第一季度销售汇总"
    ws["A1"].alignment = Alignment(horizontal="center")
    wb.save(str(path))


def correct_wide(path: Path) -> None:
    """Widen the column AND tell Excel to keep the sheet one page wide."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path))
    ws = wb["订单明细"]
    ws.column_dimensions["F"].width = 60
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    wb.save(str(path))


CORRECT_EDITS = {
    "vis_cjk.docx": correct_cjk,
    "vis_narrow.xlsx": correct_narrow,
    "vis_table.docx": correct_table,
    "vis_chart.xlsx": correct_chart,
    "vis_contrast.xlsx": correct_contrast,
    "vis_orphan.docx": correct_orphan,
    "vis_rowclip.xlsx": correct_rowclip,
    "vis_merge.xlsx": correct_merge,
    "vis_wide.xlsx": correct_wide,
}


def build(name: str, dest_dir: Path) -> Path:
    dest = dest_dir / name
    BUILDERS[name](dest)
    return dest
