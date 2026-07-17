"""Regression tests for defects found by the adversarial review."""
import pytest
from docx import Document
from openpyxl import Workbook

from office_agent.adapters.excel_adapter import ExcelAdapter
from office_agent.adapters.word_adapter import WordAdapter
from office_agent.agent.budget import BudgetTracker
from office_agent.agent.plan import Plan
from office_agent.config import Budgets
from office_agent.core.session import EditSession
from office_agent.schemas.operations import StyleParams
from office_agent.schemas.selector import ExcelSelector


# --- find_replace: replace containing find must not double-apply -------------

def test_find_replace_replace_containing_find():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("We sell one cat here.")
    run.bold = True
    result = WordAdapter.find_replace(doc, "cat", "cats")
    assert p.text == "We sell one cats here."
    assert result["count"] == 1
    assert "warning" not in result
    assert p.runs[0].bold is True  # formatting preserved


def test_find_replace_suffix_addition():
    doc = Document()
    doc.add_paragraph("Acme Inc is Inc based.")
    result = WordAdapter.find_replace(doc, "Inc", "Inc.")
    assert doc.paragraphs[0].text == "Acme Inc. is Inc. based."
    assert result["count"] == 2


# --- whole-column / whole-row ranges -----------------------------------------

def test_parse_range_whole_column_normalizes_none():
    coords = ExcelSelector(sheet="S", range="A:A").parse_range()
    assert coords[0] == 1 and coords[1] == 1  # min_row/min_col never None


def test_write_cells_whole_column_range():
    wb = Workbook()
    ws = wb.active
    coords = ExcelSelector(sheet="S", range="B:B").parse_range()
    result = ExcelAdapter.write_cells(ws, coords, [["x"], ["y"]])
    assert ws["B1"].value == "x" and ws["B2"].value == "y"
    assert result["affected"] == ["B1", "B2"]


def test_style_whole_row_range():
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "h"
    coords = ExcelSelector(sheet="S", range="1:1").parse_range()
    ExcelAdapter.apply_style(ws, coords, StyleParams(bold=True))
    assert ws["A1"].font.bold is True


# --- chart categories --------------------------------------------------------

def test_chart_first_column_becomes_categories():
    wb = Workbook()
    ws = wb.active
    for row in [["名称", "销量"], ["苹果", 120], ["香蕉", 85]]:
        ws.append(row)
    ExcelAdapter.create_chart(ws, "A1:B3", "bar")
    chart = ws._charts[0]
    assert len(chart.series) == 1  # label column is categories, not a series
    assert chart.series[0].cat is not None


# --- fail-before-mutate ------------------------------------------------------

def test_insert_table_bad_position_leaves_no_orphan():
    doc = Document()
    doc.add_paragraph("only one")
    n_tables = len(doc.tables)
    with pytest.raises(ValueError, match="out of range"):
        WordAdapter.insert_element(doc, 99, "table", [["a"]])
    assert len(doc.tables) == n_tables  # no orphaned table appended


def test_apply_style_invalid_alignment_mutates_nothing():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("text")
    with pytest.raises(ValueError, match="Invalid alignment"):
        WordAdapter.apply_style([p], StyleParams(alignment="middle", bold=True))
    assert p.runs[0].bold is None  # bold was NOT applied before the failure


def test_insert_table_at_position_inserts_before():
    doc = Document()
    doc.add_paragraph("first")
    doc.add_paragraph("second")
    WordAdapter.insert_element(doc, 1, "table", [["cell"]])
    # Table sits between 'first' and 'second' (before paragraph index 1)
    body = list(doc.element.body)
    tags = [el.tag.split("}")[-1] for el in body]
    assert tags.index("tbl") < tags.index("p") + 2  # table not at document end
    assert doc.tables[0].rows[0].cells[0].text == "cell"


# --- plan active_index -------------------------------------------------------

def test_plan_active_index_tracks_most_recent():
    plan = Plan.from_descriptions(["A", "B"])
    plan.set_status(0, "in_progress")
    assert plan.current_index == 0
    # A fails verification and stays in_progress; model moves to B.
    plan.set_status(1, "in_progress")
    assert plan.current_index == 1  # NOT hijacked by lower-index step A
    plan.set_status(1, "done")
    assert plan.current_index == 0  # falls back to still-in_progress A


# --- budget: end-turn cap and step reset -------------------------------------

def test_end_turn_repair_budget_caps():
    tracker = BudgetTracker(Budgets(max_attempts_per_step=2))
    assert not tracker.end_turn_exhausted()
    tracker.record_end_turn_repair()
    tracker.record_end_turn_repair()
    assert tracker.end_turn_exhausted()


def test_reset_steps_preserves_task_totals():
    tracker = BudgetTracker(Budgets())
    tracker.record_tool_call(0)
    tracker.record_failure(0, "sig")
    tracker.reset_steps()
    assert tracker.total_tool_calls == 1  # task counter persists
    assert tracker.step(0).attempts == 0  # per-step state reset


# --- save protection for arbitrary existing files ----------------------------

def test_save_to_refuses_unrelated_existing_file(word_doc_path, tmp_path):
    victim = tmp_path / "victim.docx"
    victim.write_bytes(word_doc_path.read_bytes())
    session = EditSession(str(word_doc_path))
    try:
        with pytest.raises(PermissionError, match="not.*created by this session"):
            session.save_to(str(victim), overwrite=False)
        # But re-saving a file we wrote ourselves is fine.
        out = tmp_path / "mine.docx"
        session.save_to(str(out))
        session.save_to(str(out))  # second save, no error
    finally:
        session.cleanup()


# --- verifier fails closed ---------------------------------------------------

def test_verifier_fails_closed_on_malformed_response():
    from office_agent.agent.verifier import verify_edit

    class TextOnlyLLM:
        def create(self, **kwargs):
            class Block:
                type = "text"
                text = "I think it looks fine!"

            class Msg:
                content = [Block()]

            return Msg()

    verdict = verify_edit(TextOnlyLLM(), "step", "ops", after_images=[])
    assert verdict.passed is False
    assert verdict.blocking
    assert any("structured verdict" in p["description"] for p in verdict.problems)


# --- per-step circuit breakers (formerly dead config) -------------------------

def test_step_tool_cap_blocks_mutating_calls(word_doc_path):
    """A step spinning on mutating calls gets a circuit-breaker error."""
    import json
    import sys
    from pathlib import Path

    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from fake_llm import FakeLLM, FakeMessage, text_block, tool_use

    from office_agent.agent.loop import AgentSession
    from office_agent.config import Config

    config = Config(api_key="fake", visual_verify=False)
    config.budgets.max_tool_calls_per_step = 2
    edit = lambda: tool_use(
        "word_edit_text",
        {
            "doc_id": "DOC",
            "selector": {"type": "paragraph", "index": 0},
            "operation": "replace",
            "text": "x",
        },
    )
    llm = FakeLLM(
        script=[
            FakeMessage(
                content=[tool_use("open_document", {"file_path": str(word_doc_path)})],
                stop_reason="tool_use",
            ),
            FakeMessage(content=[edit()], stop_reason="tool_use"),
            FakeMessage(content=[edit()], stop_reason="tool_use"),
            FakeMessage(content=[edit()], stop_reason="tool_use"),  # over cap
            FakeMessage(content=[text_block("stopping")]),
        ]
    )
    session = AgentSession(config, llm=llm)
    real_id = {}
    original = AgentSession._handle_tool

    def patched(self, name, tool_input):
        if tool_input.get("doc_id") == "DOC" and real_id:
            tool_input["doc_id"] = next(iter(real_id))
        content, is_error = original(self, name, tool_input)
        if name == "open_document" and not is_error:
            real_id[json.loads(content)["doc_id"]] = True
        return content, is_error

    import unittest.mock

    with unittest.mock.patch.object(AgentSession, "_handle_tool", patched):
        session.run_turn("edit repeatedly")
    flat = json.dumps(llm.calls, ensure_ascii=False, default=str)
    assert "without converging" in flat


# --- Word table/cell styling (was a silent no-op) ----------------------------

def test_apply_style_on_table_cell_bold_and_shading():
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Name"
    cell = t.rows[0].cells[0]
    WordAdapter.apply_style([cell], StyleParams(bold=True, bg_color="#D9D9D9"))
    assert cell.paragraphs[0].runs[0].bold is True
    from docx.oxml.ns import qn

    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd is not None and shd.get(qn("w:fill")) == "D9D9D9"


def test_apply_style_on_whole_table():
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    for row in t.rows:
        for c in row.cells:
            c.text = "x"
    result = WordAdapter.apply_style([t], StyleParams(italic=True))
    assert "table styled" in result["affected"][0]
    assert t.rows[1].cells[1].paragraphs[0].runs[0].italic is True


def test_apply_style_word_border_now_supported():
    """Borders were once unsupported; now they apply as tblBorders."""
    from docx.oxml.ns import qn

    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    result = WordAdapter.apply_style([t], StyleParams(border={"style": "single"}))
    assert "table styled" in result["affected"][0]
    assert t._tbl.tblPr.find(qn("w:tblBorders")) is not None


def test_apply_style_unsupported_target_raises_not_silent():
    class Alien:
        pass

    with pytest.raises(ValueError, match="unsupported kind"):
        WordAdapter.apply_style([Alien()], StyleParams(bold=True))


# --- word_delete_element (capability gap found by intent battery) ------------

def test_delete_paragraph():
    doc = Document()
    doc.add_paragraph("keep")
    doc.add_paragraph("remove me")
    doc.add_paragraph("keep too")
    result = WordAdapter.delete_elements([doc.paragraphs[1]])
    assert [p.text for p in doc.paragraphs] == ["keep", "keep too"]
    assert "remove me" in result["affected"][0]


def test_delete_table():
    doc = Document()
    doc.add_paragraph("text")
    doc.add_table(rows=1, cols=1)
    assert len(doc.tables) == 1
    WordAdapter.delete_elements([doc.tables[0]])
    assert len(doc.tables) == 0


def test_delete_via_tool_dispatch(word_doc_path):
    from office_agent.tools import REGISTRY, ToolContext

    ctx = ToolContext()
    try:
        opened = REGISTRY.dispatch(ctx, "open_document", {"file_path": str(word_doc_path)})
        doc_id = opened["doc_id"]
        session = ctx.sessions.get(doc_id)
        # Insert through the tool layer so the pre-delete state is snapshotted
        # (direct in-memory mutation would bypass undo history).
        inserted = REGISTRY.dispatch(
            ctx,
            "word_insert_element",
            {"doc_id": doc_id, "element_type": "paragraph", "content": "DELETE_TARGET_XYZ"},
        )
        assert inserted["success"], inserted
        result = REGISTRY.dispatch(
            ctx,
            "word_delete_element",
            {"doc_id": doc_id, "selector": {"type": "text_match", "contains": "DELETE_TARGET_XYZ"}},
        )
        assert result["success"], result
        texts = [p.text for p in session.doc.paragraphs]
        assert "DELETE_TARGET_XYZ" not in texts
        # Deletion is snapshotted -> undoable
        undo = REGISTRY.dispatch(ctx, "undo", {"doc_id": doc_id, "steps": 1})
        assert undo["success"]
        assert any("DELETE_TARGET_XYZ" in p.text for p in ctx.sessions.get(doc_id).doc.paragraphs)
    finally:
        ctx.sessions.close_all()


# --- excel_delete_chart (duplicate-chart gap found by Excel battery) ---------

def test_delete_chart_by_index():
    wb = Workbook()
    ws = wb.active
    for row in [["a", 1], ["b", 2]]:
        ws.append(row)
    ExcelAdapter.create_chart(ws, "A1:B2", "bar")
    ExcelAdapter.create_chart(ws, "A1:B2", "line")
    assert len(ws._charts) == 2
    result = ExcelAdapter.delete_chart(ws, 0)
    assert len(ws._charts) == 1
    assert "removed chart 0" in result["affected"][0]


def test_delete_all_charts():
    wb = Workbook()
    ws = wb.active
    ws.append(["a", 1])
    ExcelAdapter.create_chart(ws, "A1:B1", "bar")
    ExcelAdapter.delete_chart(ws, None)
    assert len(ws._charts) == 0


def test_delete_chart_no_charts_is_actionable():
    wb = Workbook()
    with pytest.raises(ValueError, match="no charts"):
        ExcelAdapter.delete_chart(wb.active, None)


def test_delete_chart_bad_index_is_actionable():
    wb = Workbook()
    ws = wb.active
    ws.append(["a", 1])
    ExcelAdapter.create_chart(ws, "A1:B1", "bar")
    with pytest.raises(ValueError, match="out of range"):
        ExcelAdapter.delete_chart(ws, 5)


# --- delta-review round 2 fixes ----------------------------------------------

def test_switch_signature_scoped_not_step_global():
    """A harness-breach switch must not escalate an unrelated first failure."""
    from office_agent.agent.budget import Action

    tracker = BudgetTracker(Budgets())
    assert tracker.record_failure(0, "harness:step_toolcap") == Action.RETRY
    assert tracker.record_failure(0, "harness:step_toolcap") == Action.SWITCH_STRATEGY
    # First-ever verify failure of a DIFFERENT kind: plain retry, not ask_user.
    assert tracker.record_failure(0, "verify:verify_style") == Action.RETRY
    # But the switched signature failing again does escalate.
    assert tracker.record_failure(0, "harness:step_toolcap") == Action.ASK_USER


def test_refresh_ladder_preserves_user_asks():
    tracker = BudgetTracker(Budgets())
    step = tracker.step(0)
    tracker.record_failure(0, "sig")
    tracker.record_failure(0, "sig")
    step.user_asks = 1
    step.refresh_ladder()
    assert step.attempts == 0 and step.signatures == [] and not step.strategy_switched
    assert step.user_asks == 1  # ask cap survives the refresh


def test_verifier_explicit_fail_is_always_blocking():
    from office_agent.agent.verifier import VerificationResult

    v = VerificationResult(
        passed=False,
        problems=[{"page": 0, "element_hint": "x", "description": "d", "severity": "minor"}],
        confidence=0.7,
    )
    assert v.blocking is True
    assert VerificationResult(passed=True, confidence=0.9).blocking is False
    assert VerificationResult(passed=True, skipped=True).blocking is False


def test_delete_single_cell_refused():
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    cell = t.rows[0].cells[0]
    with pytest.raises(ValueError, match="rectangular structure"):
        WordAdapter.delete_elements([cell])
    assert len(t.rows[0].cells) == 2  # table intact


def test_linked_header_edit_refused_with_owner_hint():
    from docx.enum.section import WD_SECTION

    from office_agent.core.selector_parser import SelectorError, SelectorParser

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "owner header"
    doc.add_section(WD_SECTION.NEW_PAGE)
    assert doc.sections[1].header.is_linked_to_previous
    with pytest.raises(SelectorError, match="LINKED.*section 0"):
        SelectorParser.parse_word_selector(
            {"type": "header", "section_index": 1}, doc
        )
    # Section 0 (the owner) still addressable.
    paras = SelectorParser.parse_word_selector({"type": "header"}, doc)
    assert paras and paras[0].text == "owner header"


def test_new_turn_resets_step_budgets(word_doc_path):
    import sys
    from pathlib import Path

    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from fake_llm import FakeLLM, FakeMessage, text_block

    from office_agent.agent.loop import AgentSession
    from office_agent.config import Config

    llm = FakeLLM(script=[
        FakeMessage(content=[text_block("turn1 done")]),
        FakeMessage(content=[text_block("turn2 done")]),
    ])
    session = AgentSession(Config(api_key="fake", visual_verify=False), llm=llm)
    session.run_turn("first")
    # Simulate turn 1 having consumed step budget / aged clock / abort flag.
    session.budget.step(0).tool_calls = 99
    session.budget.end_turn_repairs = 99
    session.step_started_at[0] = 1.0  # ancient
    session.abort_requested = True
    session.run_turn("second")
    assert session.budget.step(0).tool_calls == 0
    assert session.budget.end_turn_repairs == 0
    assert session.step_started_at == {}
    assert session.abort_requested is False


def test_abort_mid_batch_refuses_queued_tools(word_doc_path):
    """Once abort is set inside a multi-tool batch, queued calls are refused."""
    import json as _json
    import sys
    import unittest.mock
    from pathlib import Path

    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from fake_llm import FakeLLM, FakeMessage, text_block, tool_use

    from office_agent.agent.loop import AgentSession
    from office_agent.config import Config

    llm = FakeLLM(script=[
        FakeMessage(
            content=[
                tool_use("open_document", {"file_path": str(word_doc_path)}),
                tool_use("get_structure", {"doc_id": "whatever"}),
            ],
            stop_reason="tool_use",
        ),
        FakeMessage(content=[text_block("stopped")]),
    ])
    session = AgentSession(Config(api_key="fake", visual_verify=False), llm=llm)

    original = AgentSession._handle_tool

    def aborting_handle(self, name, tool_input):
        content, is_error = original(self, name, tool_input)
        self.abort_requested = True  # decision lands during the first call
        return content, is_error

    with unittest.mock.patch.object(AgentSession, "_handle_tool", aborting_handle):
        result = session.run_turn("do two things")
    assert result.aborted is True
    # After abort no further LLM call happens; the refusal lives in history.
    flat = _json.dumps(session.history.messages, ensure_ascii=False, default=str)
    assert "was not executed" in flat  # second tool refused, pairing intact
    tool_results = [
        b
        for m in session.history.messages
        if m["role"] == "user" and isinstance(m["content"], list)
        for b in m["content"]
        if isinstance(b, dict) and b.get("type") == "tool_result"
    ]
    assert len(tool_results) == 2  # every tool_use got its tool_result


# --- Word borders (tcBorders/tblBorders, formerly unsupported) ---------------

def test_table_borders_applied():
    from docx.oxml.ns import qn

    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    WordAdapter.apply_style(
        [t], StyleParams(border={"style": "double", "size": 1.0, "color": "#FF0000"})
    )
    borders = t._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    top = borders.find(qn("w:top"))
    assert top.get(qn("w:val")) == "double"
    assert top.get(qn("w:sz")) == "8"  # 1.0pt * 8
    assert top.get(qn("w:color")) == "FF0000"
    assert borders.find(qn("w:insideH")) is not None  # all sides by default


def test_cell_borders_specific_sides():
    from docx.oxml.ns import qn

    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    WordAdapter.apply_style(
        [cell], StyleParams(border={"style": "single", "sides": ["bottom"]})
    )
    borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
    assert borders.find(qn("w:bottom")) is not None
    assert borders.find(qn("w:top")) is None  # only the requested side


def test_border_on_paragraph_target_is_actionable():
    doc = Document()
    p = doc.add_paragraph("text")
    with pytest.raises(ValueError, match="tables, table rows, or cells"):
        WordAdapter.apply_style([p], StyleParams(border={"style": "single"}))


def test_border_invalid_style_rejected_before_mutation():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    with pytest.raises(ValueError, match="Invalid border style"):
        WordAdapter.apply_style([t], StyleParams(border={"style": "wavy"}))


def test_cell_border_rejects_inside_sides():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    with pytest.raises(ValueError, match="table-level only"):
        WordAdapter.apply_style(
            [t.rows[0].cells[0]],
            StyleParams(border={"style": "single", "sides": ["insideH"]}),
        )


def test_border_roundtrips_through_save(tmp_path):
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "x"
    WordAdapter.apply_style([t], StyleParams(border={"style": "thick", "size": 2}))
    p = tmp_path / "b.docx"
    doc.save(str(p))
    from docx.oxml.ns import qn

    reloaded = Document(str(p))
    assert reloaded.tables[0]._tbl.tblPr.find(qn("w:tblBorders")) is not None


# --- gc command --------------------------------------------------------------

def test_gc_removes_only_stale_dirs(tmp_path, monkeypatch):
    import os
    import time

    from office_agent import cli
    from office_agent.core import session as session_mod

    fake_root = tmp_path / "fallback"
    old_dir = fake_root / "old_session"
    new_dir = fake_root / "new_session"
    old_dir.mkdir(parents=True)
    new_dir.mkdir(parents=True)
    stale = time.time() - 48 * 3600
    os.utime(old_dir, (stale, stale))
    monkeypatch.setattr(cli, "console", cli.console)
    monkeypatch.setattr(session_mod, "FALLBACK_ROOT", fake_root)
    monkeypatch.setattr(session_mod, "WORD_CONTAINER", tmp_path / "nope1")
    monkeypatch.setattr(session_mod, "EXCEL_CONTAINER", tmp_path / "nope2")
    cli.gc(older_than_hours=24)
    assert not old_dir.exists()
    assert new_dir.exists()


# --- table content lenient parsing (flaky insert_table root cause) -----------

def test_insert_table_accepts_json_string_content():
    doc = Document()
    WordAdapter.insert_element(
        doc, None, "table", '[["Name","Score"],["Alice","90"]]'
    )
    t = doc.tables[-1]
    assert t.rows[0].cells[0].text == "Name"
    assert t.rows[1].cells[1].text == "90"


def test_insert_table_wraps_flat_list_as_one_row():
    doc = Document()
    result = WordAdapter.insert_element(doc, None, "table", ["a", "b", "c"])
    assert result["rows"] == 1 and result["cols"] == 3


def test_insert_table_bad_content_error_shows_example_and_got():
    doc = Document()
    with pytest.raises(ValueError, match=r"e\.g\..*Got: int"):
        WordAdapter.insert_element(doc, None, "table", 42)
