import pytest

from office_agent.core.session import EditSession, SessionManager
from office_agent.core.snapshot_manager import SnapshotManager


def _first_para_text(session):
    return session.doc.paragraphs[0].text


def test_open_creates_isolated_working_copy(word_doc_path):
    original_bytes = word_doc_path.read_bytes()
    session = EditSession(str(word_doc_path))
    try:
        assert session.working_path.exists()
        assert session.working_path != word_doc_path
        session.doc.add_paragraph("mutation")
        session.flush()
        # Original untouched even after flush
        assert word_doc_path.read_bytes() == original_bytes
    finally:
        session.cleanup()


def test_save_to_original_requires_overwrite(word_doc_path):
    session = EditSession(str(word_doc_path))
    try:
        with pytest.raises(PermissionError, match="Refusing to overwrite"):
            session.save_to(str(word_doc_path), overwrite=False)
    finally:
        session.cleanup()


def test_save_to_new_path(word_doc_path, tmp_path):
    session = EditSession(str(word_doc_path))
    try:
        session.doc.add_paragraph("saved change")
        out = session.save_to(str(tmp_path / "out.docx"))
        assert out.exists()
        from docx import Document

        assert Document(str(out)).paragraphs[-1].text == "saved change"
    finally:
        session.cleanup()


def test_unsupported_extension_raises(tmp_path):
    bad = tmp_path / "x.txt"
    bad.write_text("hi")
    with pytest.raises(ValueError, match="Unsupported file type"):
        EditSession(str(bad))


def test_session_manager_unknown_id_is_actionable():
    mgr = SessionManager()
    with pytest.raises(KeyError, match="Unknown doc_id"):
        mgr.get("nope")


def test_snapshot_undo_roundtrip(word_doc_path):
    session = EditSession(str(word_doc_path))
    try:
        snaps = SnapshotManager(session)
        before = _first_para_text(session)
        session.doc.paragraphs[0].text = "CHANGED"
        snaps.snapshot("edit_text", {"op": "replace"})
        assert _first_para_text(session) == "CHANGED"
        snaps.undo(1)
        assert _first_para_text(session) == before
    finally:
        session.cleanup()


def test_snapshot_index_survives_restart(word_doc_path):
    session = EditSession(str(word_doc_path))
    try:
        snaps = SnapshotManager(session)
        session.doc.paragraphs[0].text = "STEP1"
        snaps.snapshot("step1", {})
        # New manager instance on the same dir (simulates process restart)
        snaps2 = SnapshotManager(session)
        ids = [s["id"] for s in snaps2.list()]
        assert any("step1" in i for i in ids)
        snaps2.undo(1)
        assert _first_para_text(session) != "STEP1"
    finally:
        session.cleanup()


def test_undo_nothing_raises(word_doc_path):
    session = EditSession(str(word_doc_path))
    try:
        snaps = SnapshotManager(session)
        with pytest.raises(ValueError, match="Nothing to undo"):
            snaps.undo(1)
    finally:
        session.cleanup()


def test_restore_truncates_later_snapshots(word_doc_path):
    session = EditSession(str(word_doc_path))
    try:
        snaps = SnapshotManager(session)
        session.doc.paragraphs[0].text = "A"
        id_a = snaps.snapshot("a", {})
        session.doc.paragraphs[0].text = "B"
        snaps.snapshot("b", {})
        snaps.restore(id_a)
        assert _first_para_text(session) == "A"
        assert all("b" != s["label"] for s in snaps.list())
    finally:
        session.cleanup()
