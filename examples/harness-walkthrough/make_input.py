"""Builds the plain 'before' document for the harness walkthrough demo.

Pins an explicit East Asian font on every run. A bare `Document()`'s default
template leaves the Normal style with no font at all (verified: no <w:rPr> on
w:styleId="Normal" in styles.xml) -- opening/exporting that in Word makes Word
guess a CJK font per glyph with nothing pinned down, and its own fallback
logic inconsistently alternates between two different bundled fonts
(MS Mincho / Microsoft YaHei) within the same run, which renders as random
"bolding" in the exported PDF/PNG. Confirmed via `pdffonts` + PyMuPDF span
extraction on the pristine (never-edited) render -- this is a Word PDF-export
font-substitution artifact from this fixture script, not an office_agent bug.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

EAST_ASIAN_FONT = "Microsoft YaHei"


def set_font(run, name=EAST_ASIAN_FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


doc = Document()
title = doc.add_paragraph("季度工作总结")
set_font(title.runs[0])
title.runs[0].font.size = Pt(11)  # deliberately plain: not bold, not centered, default size

for text in [
    "本季度团队完成了三项主要工作：文档编辑 agent 的渲染验证闭环、"
    "反暴力升级梯的熔断逻辑、以及 Excel 保真度守卫。下一步计划扩展到"
    "原生自动化操作与 XML 级别的补丁能力。",
    "以下为本季度任务分配，将在下方补充为表格。",
]:
    p = doc.add_paragraph(text)
    set_font(p.runs[0])

doc.save("input/quarterly_summary.docx")
print("wrote input/quarterly_summary.docx")
