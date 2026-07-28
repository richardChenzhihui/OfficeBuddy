"""Word adapter wrapping python-docx. Stateless: mutates the elements passed in."""
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..schemas.operations import StyleParams, TextOperation

_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_east_asian_font(run: Any, name: str) -> None:
    """Sets w:rFonts/@w:eastAsia -- the attribute Word actually consults to
    pick a glyph font for CJK text. python-docx's `run.font.name` only ever
    writes @w:ascii/@w:hAnsi (the Latin-script font slots); leaving
    @w:eastAsia unset makes Word guess a CJK font per character at PDF-export
    time. Confirmed empirically: Word then inconsistently alternates between
    two different fallback fonts (e.g. MS Mincho / Microsoft YaHei) within a
    single run, which renders as seemingly-random bolding."""
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _east_asian_font(run: Any) -> Optional[str]:
    from docx.oxml.ns import qn

    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    return rfonts.get(qn("w:eastAsia")) if rfonts is not None else None


def _dominant_document_font(doc: Any) -> Optional[str]:
    """Best-effort: the first explicitly-set font anywhere in the document's
    body paragraphs, so newly-inserted content (a fresh table, a new
    paragraph) matches whatever font the rest of the document already uses
    instead of leaving @w:eastAsia unset on brand-new runs. Returns None if
    the document itself has no font pinned anywhere -- we don't invent one."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            name = _east_asian_font(run) or run.font.name
            if name:
                return name
    return None


def _copy_run_format(src: Any, dst: Any) -> None:
    ascii_name = src.font.name
    if ascii_name:
        dst.font.name = ascii_name
    east_asian_name = _east_asian_font(src) or ascii_name
    if east_asian_name:
        _set_east_asian_font(dst, east_asian_name)
    if src.font.size:
        dst.font.size = src.font.size
    dst.font.bold = src.font.bold
    dst.font.italic = src.font.italic
    dst.font.underline = src.font.underline
    if src.font.color and src.font.color.rgb:
        dst.font.color.rgb = src.font.color.rgb


_BORDER_STYLES = {"single", "double", "dashed", "dotted", "thick", "none"}
_BORDER_SIDES = {"top", "left", "bottom", "right", "insideH", "insideV"}


def _validate_border(border: dict) -> dict:
    """Normalize a border spec: {'style','size'(pt),'color','sides'}."""
    if not isinstance(border, dict):
        raise ValueError(
            "border must be a dict like {'style':'single','size':0.5,"
            "'color':'#000000','sides':['top','bottom']}."
        )
    style = border.get("style", "single")
    if style not in _BORDER_STYLES:
        raise ValueError(
            f"Invalid border style '{style}': expected one of {sorted(_BORDER_STYLES)}."
        )
    try:
        size_pt = float(border.get("size", 0.5))
    except (TypeError, ValueError):
        raise ValueError(f"Invalid border size {border.get('size')!r}: expected points, e.g. 0.5.")
    if not 0 < size_pt <= 12:
        raise ValueError(f"Border size {size_pt}pt out of range (0..12].")
    color = str(border.get("color", "#000000")).lstrip("#").upper()
    sides = border.get("sides") or []
    unknown = [s for s in sides if s not in _BORDER_SIDES]
    if unknown:
        raise ValueError(
            f"Unknown border sides {unknown}: expected {sorted(_BORDER_SIDES)}."
        )
    return {"style": style, "sz": max(2, int(size_pt * 8)), "color": color, "sides": sides}


def _make_borders_element(tag: str, border: dict, sides: list):
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    el = OxmlElement(tag)
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), border["style"])
        b.set(qn("w:sz"), str(border["sz"]))
        b.set(qn("w:color"), border["color"])
        el.append(b)
    return el


def _parse_color(value: str) -> RGBColor:
    raw = value[1:] if value.startswith("#") else value
    try:
        return RGBColor.from_string(raw)
    except Exception as exc:
        raise ValueError(
            f"Invalid color '{value}': expected hex like '#FF0000' or 'FF0000'."
        ) from exc


class WordAdapter:
    """Stateless operations on python-docx elements."""

    @staticmethod
    def read_content(elements: List[Any]) -> List[Dict[str, Any]]:
        result = []
        for element in elements:
            if hasattr(element, "rows"):  # Table
                result.append(
                    {
                        "type": "table",
                        "data": [[cell.text for cell in row.cells] for row in element.rows],
                    }
                )
            elif hasattr(element, "text"):  # Paragraph or Run
                result.append(
                    {
                        "type": "text",
                        "text": element.text,
                        "style": element.style.name if hasattr(element, "style") else None,
                    }
                )
        return result

    @staticmethod
    def edit_text(
        elements: List[Any], operation: TextOperation, text: Optional[str] = None
    ) -> Dict[str, Any]:
        affected: List[str] = []
        for element in elements:
            if not hasattr(element, "text"):
                continue
            old_text = element.text

            if operation == TextOperation.REPLACE:
                element.text = text or ""
                affected.append(f"replaced '{old_text[:50]}' -> '{(text or '')[:50]}'")

            elif operation == TextOperation.INSERT:
                if not text:
                    continue
                if hasattr(element, "runs") and element.runs:
                    # Prepend a run, preserving the paragraph's existing runs/formatting.
                    first = element.runs[0]
                    new_run = element.add_run(text)
                    _copy_run_format(first, new_run)
                    first._element.addprevious(new_run._element)
                elif hasattr(element, "runs"):
                    element.add_run(text)
                else:  # Run
                    element.text = text + element.text
                affected.append(f"inserted '{text[:50]}' at start")

            elif operation == TextOperation.DELETE:
                element.text = ""
                affected.append(f"deleted '{old_text[:50]}'")

            elif operation == TextOperation.APPEND:
                if not text:
                    continue
                if hasattr(element, "runs"):
                    if element.runs:
                        last_run = element.runs[-1]
                        new_run = element.add_run(text)
                        _copy_run_format(last_run, new_run)
                    else:
                        element.add_run(text)
                else:  # Run
                    element.text = element.text + text
                affected.append(f"appended '{text[:50]}'")

        return {"affected": affected, "operation": operation.value}

    @staticmethod
    def apply_style(elements: List[Any], style: StyleParams) -> Dict[str, Any]:
        # Validate everything BEFORE mutating: a failed call must not leave
        # half-applied styling behind.
        if style.alignment and style.alignment not in _ALIGNMENT_MAP:
            raise ValueError(
                f"Invalid alignment '{style.alignment}': "
                f"expected one of {sorted(_ALIGNMENT_MAP)}."
            )
        if style.color:
            _parse_color(style.color)
        border_spec = _validate_border(style.border) if style.border else None
        if border_spec and not any(
            hasattr(e, "rows") or hasattr(e, "cells") or hasattr(e, "_tc")
            for e in elements
        ):
            raise ValueError(
                "'border' applies only to tables, table rows, or cells — "
                "select a table / row / cell target for border styling."
            )
        affected: List[str] = []
        for element in elements:
            if hasattr(element, "rows"):  # Table: style every cell
                for row in element.rows:
                    for cell in row.cells:
                        WordAdapter._apply_cell_style(cell, style, border=None)
                if border_spec:
                    WordAdapter._apply_table_borders(element, border_spec)
                affected.append(
                    f"table styled ({len(element.rows)} rows x "
                    f"{len(element.columns)} cols)"
                )
            elif hasattr(element, "paragraphs") and hasattr(element, "_tc"):  # Cell
                WordAdapter._apply_cell_style(element, style, border=border_spec)
                affected.append("cell styled")
            elif hasattr(element, "runs"):  # Paragraph
                for run in element.runs:
                    WordAdapter._apply_run_style(run, style)
                WordAdapter._apply_paragraph_format(element, style)
                affected.append(f"paragraph styled ({len(element.runs)} runs)")
            elif hasattr(element, "font"):  # Run
                WordAdapter._apply_run_style(element, style)
                affected.append("run styled")
            elif hasattr(element, "cells"):  # Table row
                for cell in element.cells:
                    WordAdapter._apply_cell_style(cell, style, border=border_spec)
                affected.append(f"row styled ({len(element.cells)} cells)")
        if not affected:
            raise ValueError(
                "apply_style matched elements of an unsupported kind — nothing "
                "was styled. Target paragraphs, runs, tables, rows, or cells."
            )
        return {"affected": affected}

    @staticmethod
    def _apply_cell_style(cell: Any, style: StyleParams, border=None) -> None:
        for para in cell.paragraphs:
            for run in para.runs:
                WordAdapter._apply_run_style(run, style)
            WordAdapter._apply_paragraph_format(para, style)
        if style.bg_color:
            WordAdapter._shade_cell(cell, style.bg_color)
        if border:
            WordAdapter._apply_cell_borders(cell, border)

    @staticmethod
    def _apply_table_borders(table: Any, border: dict) -> None:
        from docx.oxml.ns import qn

        tbl_pr = table._tbl.tblPr
        for old in tbl_pr.findall(qn("w:tblBorders")):
            tbl_pr.remove(old)
        sides = border["sides"] or [
            "top", "left", "bottom", "right", "insideH", "insideV"
        ]
        tbl_pr.append(_make_borders_element("w:tblBorders", border, sides))

    @staticmethod
    def _apply_cell_borders(cell: Any, border: dict) -> None:
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn("w:tcBorders")):
            tc_pr.remove(old)
        sides = border["sides"] or ["top", "left", "bottom", "right"]
        bad = [s for s in sides if s.startswith("inside")]
        if bad:
            raise ValueError(
                f"Sides {bad} are table-level only; per-cell borders support "
                "top/left/bottom/right."
            )
        tc_pr.append(_make_borders_element("w:tcBorders", border, sides))

    @staticmethod
    def _shade_cell(cell: Any, hex_color: str) -> None:
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fill = hex_color.lstrip("#").upper()
        tc_pr = cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn("w:shd")):
            tc_pr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    @staticmethod
    def _apply_paragraph_format(paragraph: Any, style: StyleParams) -> None:
        fmt = paragraph.paragraph_format
        if style.paragraph_spacing is not None:
            fmt.space_before = Pt(style.paragraph_spacing)
        if style.line_spacing is not None:
            fmt.line_spacing = style.line_spacing
        if style.alignment:
            if style.alignment not in _ALIGNMENT_MAP:
                raise ValueError(
                    f"Invalid alignment '{style.alignment}': "
                    f"expected one of {sorted(_ALIGNMENT_MAP)}."
                )
            fmt.alignment = _ALIGNMENT_MAP[style.alignment]

    @staticmethod
    def _apply_run_style(run: Any, style: StyleParams) -> None:
        font = run.font
        if style.font_name:
            font.name = style.font_name
            _set_east_asian_font(run, style.font_name)
        if style.font_size:
            font.size = Pt(style.font_size)
        if style.bold is not None:
            font.bold = style.bold
        if style.italic is not None:
            font.italic = style.italic
        if style.underline is not None:
            font.underline = style.underline
        if style.color:
            font.color.rgb = _parse_color(style.color)

    @staticmethod
    def insert_element(
        doc: Document, position: Optional[int], element_type: str, content: Any
    ) -> Dict[str, Any]:
        if element_type == "table":
            # LLMs frequently serialize nested arrays as JSON strings — accept
            # that gracefully instead of failing on a formality.
            if isinstance(content, str):
                import json as _json

                try:
                    parsed = _json.loads(content)
                except (ValueError, TypeError):
                    parsed = None
                if isinstance(parsed, list):
                    content = parsed
            if isinstance(content, list) and content and not any(
                isinstance(row, list) for row in content
            ):
                # A flat list is almost certainly one intended row.
                content = [content]
            if not isinstance(content, list) or not all(
                isinstance(row, list) for row in content
            ):
                got = (
                    f"{type(content).__name__}"
                    + (
                        f" with first element {type(content[0]).__name__}"
                        if isinstance(content, list) and content
                        else ""
                    )
                )
                raise ValueError(
                    "Table content must be a 2D array of row lists, e.g. "
                    "[['Name','Score'],['Alice','90']] — pass the actual JSON "
                    f"array, not a string. Got: {got}."
                )
            # Validate position BEFORE mutating: a failed call must leave no
            # orphaned table behind.
            if position is not None and not 0 <= position < len(doc.paragraphs):
                raise ValueError(
                    f"Insert position {position} out of range: "
                    f"document has {len(doc.paragraphs)} paragraphs."
                )
            cols = len(content[0]) if content else 0
            table = doc.add_table(rows=len(content), cols=cols)
            try:
                table.style = "Table Grid"  # visible borders by default
            except KeyError:
                pass  # style absent from this document's template
            # New cells get a fresh run with no font at all; match whatever
            # font the rest of the document already uses so the new table
            # doesn't get Word's inconsistent per-glyph CJK font fallback.
            font_name = _dominant_document_font(doc)
            for i, row_data in enumerate(content):
                for j, cell_data in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = str(cell_data)
                    if font_name and cell.paragraphs[0].runs:
                        run = cell.paragraphs[0].runs[0]
                        run.font.name = run.font.name or font_name
                        _set_east_asian_font(run, font_name)
            if position is not None:
                # Insert BEFORE the position paragraph — same semantics as
                # element_type='paragraph'.
                doc.paragraphs[position]._element.addprevious(table._element)
            return {"element_type": "table", "rows": len(content), "cols": cols}

        if element_type == "paragraph":
            if position is not None:
                if not 0 <= position < len(doc.paragraphs):
                    raise ValueError(
                        f"Insert position {position} out of range: "
                        f"document has {len(doc.paragraphs)} paragraphs."
                    )
                new_para = doc.paragraphs[position].insert_paragraph_before(str(content))
            else:
                new_para = doc.add_paragraph(str(content))
            font_name = _dominant_document_font(doc)
            if font_name and new_para.runs:
                run = new_para.runs[0]
                run.font.name = run.font.name or font_name
                _set_east_asian_font(run, font_name)
            return {"element_type": "paragraph", "text": str(content)[:50]}

        if element_type == "page_break":
            from docx.enum.text import WD_BREAK

            if position is not None and 0 <= position < len(doc.paragraphs):
                para = doc.paragraphs[position]
            else:
                para = doc.add_paragraph()
            para.add_run().add_break(WD_BREAK.PAGE)
            return {"element_type": "page_break"}

        raise ValueError(
            f"Unsupported element type '{element_type}': "
            "expected table, paragraph, or page_break."
        )

    @staticmethod
    def delete_elements(elements: List[Any]) -> Dict[str, Any]:
        """Remove paragraphs/tables/rows from the document tree entirely."""
        removed: List[str] = []
        for element in elements:
            if hasattr(element, "_tc"):  # single table cell
                raise ValueError(
                    "Deleting a single table CELL would corrupt the table's "
                    "rectangular structure. Delete the whole row or table "
                    "instead, or clear the cell's text with word_edit_text "
                    "(operation='delete')."
                )
            xml_el = getattr(element, "_element", None)
            if xml_el is None or xml_el.getparent() is None:
                raise ValueError(
                    "This element kind cannot be deleted (or is already "
                    "detached). Deletable: paragraphs, tables, table rows."
                )
            if hasattr(element, "rows"):
                label = f"table({len(element.rows)}x{len(element.columns)})"
            elif hasattr(element, "text"):
                label = f"'{element.text[:50]}'"
            else:
                label = type(element).__name__
            xml_el.getparent().remove(xml_el)
            removed.append(label)
        return {"affected": removed}

    @staticmethod
    def find_replace(
        doc: Document, find: str, replace: str, scope: Optional[List[Any]] = None
    ) -> Dict[str, Any]:
        if not find:
            raise ValueError("find string must be non-empty.")
        count = 0
        affected: List[str] = []
        formatting_lost_in: List[str] = []
        elements = scope if scope else doc.paragraphs

        for element in elements:
            if not hasattr(element, "text"):
                continue
            original_text = element.text
            occurrences = original_text.count(find)
            if not occurrences:
                continue
            if hasattr(element, "runs"):
                # Decide run-level vs paragraph-level from PRE-mutation counts:
                # re-scanning after replacement misfires whenever `replace`
                # itself contains `find` (e.g. cat -> cats).
                run_occurrences = sum(run.text.count(find) for run in element.runs)
                if run_occurrences == occurrences:
                    # Every occurrence sits inside a single run — formatting safe.
                    for run in element.runs:
                        if find in run.text:
                            run.text = run.text.replace(find, replace)
                else:
                    # Some occurrence spans a run boundary: paragraph-level
                    # fallback, which flattens the paragraph's run formatting.
                    element.text = original_text.replace(find, replace)
                    formatting_lost_in.append(element.text[:50])
            else:
                element.text = original_text.replace(find, replace)
            count += occurrences
            affected.append(f"'{element.text[:50]}'")

        result: Dict[str, Any] = {
            "count": count,
            "affected": affected,
            "find": find,
            "replace": replace,
        }
        if formatting_lost_in:
            result["warning"] = (
                "Some matches spanned formatting boundaries; run formatting was "
                f"flattened in {len(formatting_lost_in)} paragraph(s)."
            )
        return result
