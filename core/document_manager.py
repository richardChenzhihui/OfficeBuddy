"""Document manager for lifecycle management"""
import os
import uuid
from typing import Dict, Optional, Any
from pathlib import Path
from enum import Enum

from docx import Document as WordDocument
from openpyxl import load_workbook, Workbook

from .snapshot_manager import SnapshotManager
from .selector_parser import SelectorParser


class DocumentType(str, Enum):
    WORD = "word"
    EXCEL = "excel"


class DocumentManager:
    """Manages document lifecycle: open, modify, save, snapshot"""
    
    def __init__(self, snapshot_dir: str = ".snapshots"):
        self.documents: Dict[str, Dict[str, Any]] = {}
        self.snapshot_manager = SnapshotManager(snapshot_dir)
        self.selector_parser = SelectorParser()
    
    def open_document(self, file_path: str) -> Dict[str, Any]:
        """Open a document and return doc_id and summary"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = os.path.abspath(file_path)
        file_ext = Path(file_path).suffix.lower()
        
        doc_id = str(uuid.uuid4())
        
        if file_ext == ".docx":
            doc = WordDocument(file_path)
            doc_type = DocumentType.WORD
            summary = {
                "type": "word",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
        elif file_ext in [".xlsx", ".xlsm"]:
            doc = load_workbook(file_path)
            doc_type = DocumentType.EXCEL
            summary = {
                "type": "excel",
                "sheets": doc.sheetnames,
                "sheet_count": len(doc.sheetnames)
            }
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Store document
        self.documents[doc_id] = {
            "doc_id": doc_id,
            "file_path": file_path,
            "doc_type": doc_type.value,
            "document": doc,
            "previews": {},  # preview_id -> preview_data
            "modified": False
        }
        
        return {
            "doc_id": doc_id,
            "summary": summary
        }
    
    def get_structure(self, doc_id: str, depth: int = 2) -> Dict[str, Any]:
        """Get document structure tree"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        
        doc_info = self.documents[doc_id]
        doc = doc_info["document"]
        doc_type = doc_info["doc_type"]
        
        if doc_type == "word":
            structure = {
                "type": "word",
                "paragraphs": [
                    {
                        "index": i,
                        "text": para.text[:100] if len(para.text) > 100 else para.text,
                        "runs": len(para.runs),
                        "style": para.style.name if para.style else None
                    }
                    for i, para in enumerate(doc.paragraphs)
                ],
                "tables": [
                    {
                        "index": i,
                        "rows": len(table.rows),
                        "cols": len(table.columns) if table.rows else 0
                    }
                    for i, table in enumerate(doc.tables)
                ]
            }
        else:  # excel
            structure = {
                "type": "excel",
                "sheets": []
            }
            for sheet_name in doc.sheetnames:
                sheet = doc[sheet_name]
                structure["sheets"].append({
                    "name": sheet_name,
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                    "dimensions": sheet.dimensions
                })
        
        return structure
    
    def get_document(self, doc_id: str) -> Any:
        """Get document object"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        return self.documents[doc_id]["document"]
    
    def get_doc_info(self, doc_id: str) -> Dict[str, Any]:
        """Get document info"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        return self.documents[doc_id]
    
    def create_preview(self, doc_id: str, preview_id: str, preview_data: Dict[str, Any]):
        """Store preview data"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        self.documents[doc_id]["previews"][preview_id] = preview_data
    
    def get_preview(self, doc_id: str, preview_id: str) -> Optional[Dict[str, Any]]:
        """Get preview data"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        return self.documents[doc_id]["previews"].get(preview_id)
    
    def apply_changes(self, doc_id: str, preview_id: str) -> Dict[str, Any]:
        """Apply previewed changes and create snapshot"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        
        preview = self.get_preview(doc_id, preview_id)
        if not preview:
            raise ValueError(f"Preview not found: {preview_id}")
        
        doc_info = self.documents[doc_id]
        file_path = doc_info["file_path"]
        doc = doc_info["document"]
        doc_type = doc_info["doc_type"]
        
        # Create snapshot before applying changes
        snapshot_id = self.snapshot_manager.create_snapshot(doc_id, file_path)
        
        # Apply changes based on operation type
        operation = preview.get("operation")
        # Merge result and other preview data
        preview_data = {**preview}
        preview_data.pop("operation", None)  # Remove operation from data
        
        if doc_type == "excel":
            from ..adapters.excel_adapter import ExcelAdapter
            from ..core.selector_parser import SelectorParser
            from ..schemas.selector import ExcelSelector
            
            if operation == "excel_write_cells":
                selector = ExcelSelector(**preview_data.get("selector", {}))
                worksheet, coords = self.selector_parser.parse_excel_selector(selector, doc)
                if worksheet:
                    from ..schemas.operations import FillMode
                    fill_mode = FillMode(preview_data.get("fill_mode", "overwrite"))
                    ExcelAdapter.write_cells(
                        worksheet, coords, 
                        preview_data.get("values", []),
                        fill_mode
                    )
            
            elif operation == "excel_edit_formula":
                selector = ExcelSelector(**preview_data.get("selector", {}))
                worksheet, coords = self.selector_parser.parse_excel_selector(selector, doc)
                if worksheet:
                    ExcelAdapter.edit_formula(worksheet, coords, preview_data.get("formula", ""))
            
            elif operation == "excel_edit_style":
                from ..schemas.operations import StyleParams
                selector = ExcelSelector(**preview_data.get("selector", {}))
                worksheet, coords = self.selector_parser.parse_excel_selector(selector, doc)
                if worksheet:
                    style = StyleParams(**preview_data.get("style_params", {}))
                    ExcelAdapter.apply_style(worksheet, coords, style)
            
            elif operation == "excel_insert_rows_cols":
                sheet_name = preview_data.get("sheet")
                if sheet_name in doc.sheetnames:
                    worksheet = doc[sheet_name]
                    ExcelAdapter.insert_rows_cols(
                        worksheet,
                        preview_data.get("position", 1),
                        preview_data.get("count", 1),
                        preview_data.get("insert_type", "row")
                    )
            
            elif operation == "excel_create_chart":
                sheet_name = preview_data.get("sheet")
                if sheet_name in doc.sheetnames:
                    worksheet = doc[sheet_name]
                    ExcelAdapter.create_chart(
                        worksheet,
                        preview_data.get("data_range", ""),
                        preview_data.get("chart_type", "bar"),
                        preview_data.get("chart_options")
                    )
        
        elif doc_type == "word":
            from ..adapters.word_adapter import WordAdapter
            from ..core.selector_parser import SelectorParser
            from ..schemas.selector import WordSelector
            from ..schemas.operations import TextOperation, StyleParams
            
            if operation == "word_edit_text":
                selector = WordSelector(**preview_data.get("selector", {}))
                elements = self.selector_parser.parse_word_selector(selector, doc)
                if elements:
                    op = TextOperation(preview_data.get("operation", "replace"))
                    # Re-apply the edit (preview already calculated, now actually modify)
                    WordAdapter.edit_text(elements, op, preview_data.get("text"))
            
            elif operation == "word_edit_style":
                selector = WordSelector(**preview_data.get("selector", {}))
                elements = self.selector_parser.parse_word_selector(selector, doc)
                if elements:
                    style = StyleParams(**preview_data.get("style_params", {}))
                    WordAdapter.apply_style(elements, style)
            
            elif operation == "word_insert_element":
                WordAdapter.insert_element(
                    doc,
                    preview_data.get("position"),
                    preview_data.get("element_type", "paragraph"),
                    preview_data.get("content")
                )
            
            elif operation == "word_find_replace":
                scope = preview_data.get("scope")
                elements = None
                if scope:
                    selector = WordSelector(**scope)
                    elements = self.selector_parser.parse_word_selector(selector, doc)
                WordAdapter.find_replace(doc, preview_data.get("find", ""), 
                                        preview_data.get("replace", ""), elements)
        
        doc_info["modified"] = True
        
        return {
            "success": True,
            "snapshot_id": snapshot_id
        }
    
    def save_document(self, doc_id: str, path: Optional[str] = None) -> bool:
        """Save document"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        
        doc_info = self.documents[doc_id]
        doc = doc_info["document"]
        save_path = path or doc_info["file_path"]
        
        if doc_info["doc_type"] == "word":
            doc.save(save_path)
        else:  # excel
            doc.save(save_path)
        
        doc_info["modified"] = False
        return True
    
    def undo(self, doc_id: str) -> bool:
        """Undo last change"""
        snapshot_id = self.snapshot_manager.get_latest_snapshot(doc_id)
        if not snapshot_id:
            return False
        
        return self.restore_snapshot(doc_id, snapshot_id)
    
    def restore_snapshot(self, doc_id: str, snapshot_id: str) -> bool:
        """Restore document from snapshot"""
        if doc_id not in self.documents:
            raise ValueError(f"Document not found: {doc_id}")
        
        doc_info = self.documents[doc_id]
        file_path = doc_info["file_path"]
        
        success = self.snapshot_manager.restore_snapshot(doc_id, snapshot_id, file_path)
        if success:
            # Reload document
            if doc_info["doc_type"] == "word":
                doc_info["document"] = WordDocument(file_path)
            else:
                doc_info["document"] = load_workbook(file_path)
            doc_info["modified"] = False
        
        return success
