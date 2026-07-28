import pytest
from docx import Document

from office_agent.adapters.word_adapter import WordAdapter
from office_agent.schemas.operations import StyleParams, TextOperation


@pytest.fixture
def doc():
    d = Document()
    p = d.add_paragraph()
    run = p.add_run("Hello ")
    run.bold = True
    p.add_run("world")
    d.add_paragraph("Second paragraph")
    return d


def test_append_preserves_existing_runs_and_adds_once(doc):
    para = doc.paragraphs[0]
    WordAdapter.edit_text([para], TextOperation.APPEND, " suffix")
    assert para.text == "Hello world suffix"
    assert para.runs[0].bold is True  # existing formatting untouched


def test_append_copies_last_run_format(doc):
    para = doc.paragraphs[0]
    para.runs[-1].italic = True
    WordAdapter.edit_text([para], TextOperation.APPEND, "!")
    assert para.runs[-1].italic is True


def test_insert_prepends_preserving_runs(doc):
    para = doc.paragraphs[0]
    WordAdapter.edit_text([para], TextOperation.INSERT, "Start: ")
    assert para.text == "Start: Hello world"
    # Original bold run still present
    assert any(r.bold for r in para.runs)


def test_replace(doc):
    para = doc.paragraphs[1]
    WordAdapter.edit_text([para], TextOperation.REPLACE, "New text")
    assert para.text == "New text"


def test_delete(doc):
    para = doc.paragraphs[1]
    WordAdapter.edit_text([para], TextOperation.DELETE)
    assert para.text == ""


def test_apply_style_sets_alignment_on_paragraph(doc):
    """Regression: old code's paragraph_format branch was unreachable."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    para = doc.paragraphs[0]
    WordAdapter.apply_style([para], StyleParams(alignment="center", font_size=14.0))
    assert para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert all(r.font.size.pt == 14.0 for r in para.runs)


def test_apply_style_invalid_alignment_raises(doc):
    with pytest.raises(ValueError, match="Invalid alignment"):
        WordAdapter.apply_style([doc.paragraphs[0]], StyleParams(alignment="middle"))


def test_apply_style_invalid_color_raises(doc):
    with pytest.raises(ValueError, match="Invalid color"):
        WordAdapter.apply_style([doc.paragraphs[0]], StyleParams(color="not-a-color"))


def test_find_replace_counts_actual_occurrences(doc):
    doc.add_paragraph("aaa bbb aaa")
    result = WordAdapter.find_replace(doc, "aaa", "xxx")
    assert result["count"] == 2
    assert doc.paragraphs[-1].text == "xxx bbb xxx"


def test_find_replace_preserves_run_formatting_within_run(doc):
    para = doc.paragraphs[0]  # "Hello " (bold) + "world"
    WordAdapter.find_replace(doc, "world", "there")
    assert para.text == "Hello there"
    assert para.runs[0].bold is True
    assert "warning" not in WordAdapter.find_replace(doc, "there", "world")


def test_find_replace_empty_find_raises(doc):
    with pytest.raises(ValueError, match="non-empty"):
        WordAdapter.find_replace(doc, "", "x")


def test_insert_element_paragraph_at_position(doc):
    WordAdapter.insert_element(doc, 1, "paragraph", "Inserted")
    assert doc.paragraphs[1].text == "Inserted"


def test_insert_element_bad_position_raises(doc):
    with pytest.raises(ValueError, match="out of range"):
        WordAdapter.insert_element(doc, 99, "paragraph", "x")


def test_insert_element_table(doc):
    result = WordAdapter.insert_element(doc, None, "table", [["a", "b"], ["c", "d"]])
    assert result["rows"] == 2
    assert doc.tables[-1].rows[0].cells[0].text == "a"


def test_insert_element_bad_type_raises(doc):
    with pytest.raises(ValueError, match="Unsupported element type"):
        WordAdapter.insert_element(doc, None, "image", "x")


def _east_asian_font(run):
    from docx.oxml.ns import qn

    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    return rfonts.get(qn("w:eastAsia")) if rfonts is not None else None


def test_apply_style_font_name_sets_east_asian_font(doc):
    """Regression: python-docx's font.name only ever writes w:ascii/w:hAnsi;
    CJK glyphs render via the separate w:eastAsia attribute, so a requested
    font change would otherwise be silently invisible for Chinese text."""
    WordAdapter.apply_style([doc.paragraphs[0]], StyleParams(font_name="SimSun"))
    assert all(_east_asian_font(r) == "SimSun" for r in doc.paragraphs[0].runs)


def test_insert_preserves_east_asian_font(doc):
    para = doc.paragraphs[0]
    WordAdapter.apply_style([para], StyleParams(font_name="SimSun"))
    WordAdapter.edit_text([para], TextOperation.INSERT, "前缀")
    assert _east_asian_font(para.runs[0]) == "SimSun"  # the prepended run


def test_insert_element_table_inherits_document_font(doc):
    """New table cells get a fresh run with no font at all; they must match
    whatever font the rest of the document already uses, or Word's own
    per-glyph CJK font fallback renders as seemingly-random bolding."""
    WordAdapter.apply_style([doc.paragraphs[0]], StyleParams(font_name="SimSun"))
    WordAdapter.insert_element(doc, None, "table", [["表头"]])
    cell_run = doc.tables[-1].rows[0].cells[0].paragraphs[0].runs[0]
    assert _east_asian_font(cell_run) == "SimSun"


def test_insert_element_table_no_font_invented_when_document_has_none(doc):
    """If nothing in the document has an explicit font, don't invent one."""
    WordAdapter.insert_element(doc, None, "table", [["x"]])
    cell_run = doc.tables[-1].rows[0].cells[0].paragraphs[0].runs[0]
    assert _east_asian_font(cell_run) is None


def test_insert_element_paragraph_inherits_document_font(doc):
    WordAdapter.apply_style([doc.paragraphs[0]], StyleParams(font_name="SimSun"))
    WordAdapter.insert_element(doc, None, "paragraph", "新段落")
    assert _east_asian_font(doc.paragraphs[-1].runs[0]) == "SimSun"
